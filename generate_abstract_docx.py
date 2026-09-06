#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
USSF 2026 - Academic Abstract DOCX Generator
--------------------------------------------
Generates standardized Microsoft Word documents (.docx) strictly following the
official Bogomolets NMU surgical forum template:
- Page format: A4 (210 x 297 mm)
- Margins: 25.4 mm (1 inch) all around (Top, Bottom, Left, Right)
- Base font: Times New Roman, 12 pt, black
- Line spacing: 1.0 (Single) throughout the entire document, 0 pt before/after
- Title: UPPERCASE, Bold, Centered
- Author: Surname + Initials (e.g. 'Іванченко І. І.'), Italic, Centered
- Affiliation: Italic, Left-aligned
- Section headings: Bold ('Вступ: ', 'Мета: ', 'Матеріали і методи: ', 'Результати: ', 'Висновок: ', 'Ключові слова: ')
- Paragraphs preserved with 10 mm first line indent, continuous 1.0 spacing
- References: 'Джерела:' heading with itemized sources
"""

import os
import sys
import json
import re
import argparse
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

import docx
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def format_author_initials(full_name: str) -> str:
    """
    Converts full name to Ukrainian academic standard: Surname + Initials.
    Examples:
      'Іванченко Іван Іванович' -> 'Іванченко І. І.'
      'Кузнецов Олександр Олександрович' -> 'Кузнецов О. О.'
      'Кузнецов О. О.' -> 'Кузнецов О. О.'
      'Кузнецов О.О.' -> 'Кузнецов О. О.'
      'Smith John Michael' -> 'Smith J. M.'
      'Петренко Петро' -> 'Петренко П.'
    """
    if not full_name:
        return "Автор"
    parts = full_name.strip().split()
    if not parts:
        return "Автор"
    
    last_name = parts[0][0].upper() + parts[0][1:] if parts[0] else "Автор"
    initials = []
    
    for p in parts[1:]:
        clean_p = p.replace('.', '').replace(',', '').strip()
        if not clean_p:
            continue
        if len(clean_p) <= 2 and clean_p.isupper():
            for char in clean_p:
                initials.append(f"{char.upper()}.")
        else:
            initials.append(f"{clean_p[0].upper()}.")
            
    if initials:
        return f"{last_name} {' '.join(initials)}"
    return last_name


def clean_section_text(text: str, default_label: str, aliases: list = None) -> str:
    """Removes duplicate typed section prefixes, strips tabs and ensures first letter is capitalized."""
    if not text:
        return ""
    cleaned = text.strip()
    patterns = [default_label] + (aliases or [])
    for pat in patterns:
        reg_punct = re.compile(r'^' + re.escape(pat) + r'\s*[:\-–—]\s*', re.IGNORECASE)
        if reg_punct.search(cleaned):
            cleaned = reg_punct.sub('', cleaned).strip()
            break
        reg_exact_line = re.compile(r'^' + re.escape(pat) + r'\s*(?:\r?\n|$)\s*', re.IGNORECASE)
        if reg_exact_line.search(cleaned):
            cleaned = reg_exact_line.sub('', cleaned).strip()
            break
    cleaned = cleaned.lstrip('\t ').strip()
    if cleaned:
        cleaned = cleaned[0].upper() + cleaned[1:]
    return cleaned


def add_blank_line(doc):
    """Adds a blank paragraph with 1.0 line spacing and 0 pt margins."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def create_abstract_docx(data: dict, output_path: str = None) -> str:
    """
    Builds an official DOCX abstract conforming to NMU surgical forum guidelines.
    """
    base_dir = os.path.dirname(os.path.abspath(__file__))
    submissions_dir = os.path.join(base_dir, 'заявки_тези')
    os.makedirs(submissions_dir, exist_ok=True)

    full_name = data.get('fullName', '').strip() or data.get('full_name', '').strip() or f"{data.get('last_name', '')} {data.get('first_name', '')} {data.get('middle_name', '')}".strip() or 'Учасник'
    author_initials = format_author_initials(full_name)
    title = (data.get('abstractTitle', '').strip() or data.get('title', '').strip() or 'ТЕМА НАУКОВОЇ РОБОТИ').upper()

    scientific_supervisor = data.get('scientificSupervisor', '').strip() or data.get('scientific_advisor', '').strip() or data.get('advisor', '').strip()
    department = data.get('department', '').strip()
    head_of_department = data.get('headOfDepartment', '').strip() or data.get('head_of_department', '').strip()
    institution = data.get('institution', '').strip() or data.get('affiliation', '').strip() or 'Національний медичний університет імені О. О. Богомольця'
    city_country = data.get('cityCountry', '').strip() or data.get('city_country', '').strip() or 'м. Київ, Україна'

    intro = data.get('abstractIntro', '').strip() or data.get('introduction', '').strip() or data.get('intro', '').strip()
    aim = data.get('abstractAim', '').strip() or data.get('aim', '').strip()
    materials = data.get('abstractMaterials', '').strip() or data.get('methods', '').strip() or data.get('materials', '').strip()
    results = data.get('abstractResults', '').strip() or data.get('results', '').strip() or data.get('abstractBody', '').strip()
    conclusion = data.get('abstractConclusion', '').strip() or data.get('conclusion', '').strip() or data.get('conclusions', '').strip()
    keywords = data.get('abstractKeywords', '').strip() or data.get('keywords', '').strip()
    references = data.get('abstractReferences', '').strip() or data.get('references', '').strip()

    # Determine output path
    if not output_path:
        timestamp_slug = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
        safe_name = re.sub(r'[^\w]+', '_', author_initials.replace('.', '').strip()).strip('_') or 'Учасник'
        filename = f"Тези_{safe_name}_{timestamp_slug}.docx"
        output_path = os.path.join(submissions_dir, filename)

    # Initialize Document
    template_path = os.path.join(base_dir, 'template_base.docx')
    has_template = os.path.exists(template_path)
    if has_template:
        doc = docx.Document(template_path)
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
    else:
        doc = docx.Document()

    # Configure Page Setup (A4, 25.4 mm / 1 inch margins all around)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)

    # Base style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing = 1.0
    style_normal.paragraph_format.space_before = Pt(0)
    style_normal.paragraph_format.space_after = Pt(0)

    # 1. Title (Uppercase, Bold, Centered, 12pt, 1.0 spacing)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 1.0
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    r_title = p_title.add_run(title)
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(12)

    # Blank line after title
    add_blank_line(doc)

    # 2. Author (Surname + initials, Italic, Centered, 12pt, 1.0 spacing)
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.line_spacing = 1.0
    p_author.paragraph_format.space_before = Pt(0)
    p_author.paragraph_format.space_after = Pt(0)
    r_author = p_author.add_run(author_initials)
    r_author.italic = True
    r_author.font.name = 'Times New Roman'
    r_author.font.size = Pt(12)

    # Blank line after author
    add_blank_line(doc)

    # 3. Affiliation block (Left-aligned, Italic, 12pt, 1.0 spacing)
    affil_items = []
    if scientific_supervisor:
        prefix = "" if scientific_supervisor.lower().startswith("науковий керівник") else "Науковий керівник: "
        affil_items.append(f"{prefix}{scientific_supervisor}")
    if department:
        prefix = "" if department.lower().startswith("кафедра") else "Кафедра "
        affil_items.append(f"{prefix}{department}")
    if head_of_department:
        prefix = "" if head_of_department.lower().startswith("завідувач кафедри") else "Завідувач кафедри: "
        affil_items.append(f"{prefix}{head_of_department}")
    if institution:
        affil_items.append(institution)
    if city_country:
        affil_items.append(city_country)

    for item in affil_items:
        p_aff = doc.add_paragraph()
        p_aff.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_aff.paragraph_format.line_spacing = 1.0
        p_aff.paragraph_format.space_before = Pt(0)
        p_aff.paragraph_format.space_after = Pt(0)
        r_aff = p_aff.add_run(item)
        r_aff.italic = True
        r_aff.font.name = 'Times New Roman'
        r_aff.font.size = Pt(12)

    # Blank line after affiliation
    add_blank_line(doc)

    # 4. Structured Scientific Sections
    # Each section preserves user-entered paragraphs with 10 mm indent.
    # No extra gaps between sections: continuous 1.0 line spacing.
    sections_data = [
        ("Вступ:", clean_section_text(intro, "Вступ", ["Вступ:", "Актуальність", "Актуальність:"])),
        ("Мета:", clean_section_text(aim, "Мета", ["Мета:", "Мета роботи", "Мета роботи:"])),
        ("Матеріали і методи:", clean_section_text(materials, "Матеріали і методи", ["Матеріали та методи", "Методи дослідження", "Методи дослідження:"])),
        ("Результати:", clean_section_text(results, "Результати", ["Результати:"])),
        ("Висновок:", clean_section_text(conclusion, "Висновок", ["Висновки", "Висновок:", "Висновки:"])),
        ("Ключові слова:", clean_section_text(keywords, "Ключові слова", ["Ключові слова:"]))
    ]

    for heading, sec_text in sections_data:
        if not sec_text:
            continue
        
        # Split text by user paragraphs while keeping clean strings
        raw_paras = [p.lstrip('\t ').strip() for p in sec_text.split('\n') if p.strip()]
        if not raw_paras:
            continue

        # Ensure first paragraph starts with uppercase
        p0_text = raw_paras[0]
        if p0_text:
            p0_text = p0_text[0].upper() + p0_text[1:]

        # First paragraph includes bold heading prefix
        p_first = doc.add_paragraph()
        p_first.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_first.paragraph_format.first_line_indent = Mm(10)
        p_first.paragraph_format.line_spacing = 1.0
        p_first.paragraph_format.space_before = Pt(0)
        p_first.paragraph_format.space_after = Pt(0)

        r_head = p_first.add_run(f"{heading} ")
        r_head.bold = True
        r_head.font.name = 'Times New Roman'
        r_head.font.size = Pt(12)

        r_body = p_first.add_run(p0_text)
        r_body.font.name = 'Times New Roman'
        r_body.font.size = Pt(12)

        # Subsequent paragraphs in the same section
        for subsequent_text in raw_paras[1:]:
            sub_text = subsequent_text.lstrip('\t ').strip()
            if sub_text:
                sub_text = sub_text[0].upper() + sub_text[1:]
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_sub.paragraph_format.first_line_indent = Mm(10)
            p_sub.paragraph_format.line_spacing = 1.0
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after = Pt(0)

            r_sub = p_sub.add_run(sub_text)
            r_sub.font.name = 'Times New Roman'
            r_sub.font.size = Pt(12)

    # 5. References / Джерела
    if references:
        # Blank line before references
        add_blank_line(doc)

        p_ref_head = doc.add_paragraph()
        p_ref_head.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_ref_head.paragraph_format.first_line_indent = Mm(10)
        p_ref_head.paragraph_format.line_spacing = 1.0
        p_ref_head.paragraph_format.space_before = Pt(0)
        p_ref_head.paragraph_format.space_after = Pt(0)

        r_ref_head = p_ref_head.add_run("Джерела:")
        r_ref_head.bold = True
        r_ref_head.font.name = 'Times New Roman'
        r_ref_head.font.size = Pt(12)

        ref_items = [r.lstrip('\t ').strip() for r in references.split('\n') if r.strip()]
        for idx, ref_item in enumerate(ref_items):
            clean_item = re.sub(r'^(\[\d+\]|\d+[\.\)\s\t]+)', '', ref_item).strip()
            clean_item = clean_item.lstrip('\t ').strip()
            if not clean_item:
                continue
            clean_item = clean_item[0].upper() + clean_item[1:]

            p_ref = doc.add_paragraph()
            p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_ref.paragraph_format.line_spacing = 1.0
            p_ref.paragraph_format.space_before = Pt(0)
            p_ref.paragraph_format.space_after = Pt(0)

            if has_template:
                numPr = OxmlElement('w:numPr')
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'), '0')
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'), '1')
                numPr.append(ilvl)
                numPr.append(numId)
                p_ref._p.get_or_add_pPr().append(numPr)

                r_ref = p_ref.add_run(clean_item)
                r_ref.font.name = 'Times New Roman'
                r_ref.font.size = Pt(12)
            else:
                p_ref.paragraph_format.left_indent = Mm(10)
                p_ref.paragraph_format.first_line_indent = Mm(-5)
                r_ref = p_ref.add_run(f"{idx + 1}.\t{clean_item}")
                r_ref.font.name = 'Times New Roman'
                r_ref.font.size = Pt(12)

    # Save generated DOCX file
    doc.save(output_path)
    return os.path.abspath(output_path)


def load_email_config() -> dict:
    """Load email configuration from email_config.json or environment variables."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    config_file = os.path.join(base_dir, 'email_config.json')
    config = {
        'smtp_host': os.environ.get('SMTP_HOST', 'smtp.gmail.com'),
        'smtp_port': int(os.environ.get('SMTP_PORT', 587)),
        'smtp_user': os.environ.get('SMTP_USER', ''),
        'smtp_pass': os.environ.get('SMTP_PASS', ''),
        'committee_email': os.environ.get('COMMITTEE_EMAIL', 'derk.boryslav@gmail.com'),
        'send_copy_to_author': True,
        'sync_secret_token': os.environ.get('SYNC_SECRET_TOKEN', 'ussf_secure_sync_2026_med_nmu'),
        'program_pdf_path': os.environ.get('PROGRAM_PDF_PATH', 'USSF2026_Program.pdf'),
        'google_sheet_webhook_url': os.environ.get('GOOGLE_SHEET_WEBHOOK_URL', '')
    }
    if os.path.exists(config_file):
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
                config.update({k: v for k, v in loaded.items() if v not in (None, '') or k in ('smtp_user', 'smtp_pass')})
        except Exception as e:
            print(f"[WARN] Failed to read email_config.json: {e}")
    return config


def send_abstract_email_docx(docx_path: str, data: dict, recipient: str = None) -> dict:
    """
    Sends dual notification emails with the compiled .docx abstract attached:
    1. Automated confirmation to the participant with Program PDF and compiled Word document.
    2. Committee notification to derk.boryslav@gmail.com with complete application data.
    """
    cfg = load_email_config()
    base_dir = os.path.dirname(os.path.abspath(__file__))

    smtp_host = cfg.get('smtp_host', 'smtp.gmail.com')
    smtp_port = int(cfg.get('smtp_port', 587))
    smtp_user = cfg.get('smtp_user', '').strip()
    smtp_pass = cfg.get('smtp_pass', '').strip()

    committee_email = recipient or cfg.get('committee_email', 'derk.boryslav@gmail.com')
    author_email = data.get('email', '').strip()

    if not smtp_user or not smtp_pass:
        return {
            "sent": False,
            "error": "SMTP_NOT_CONFIGURED",
            "message": "Поштові реквізити не налаштовано у файлі email_config.json (потрібно вказати smtp_user та smtp_pass).",
            "smtp_user": smtp_user,
            "committee_email": committee_email
        }

    full_name = data.get('fullName', '').strip() or 'Учасник'
    author_initials = format_author_initials(full_name)
    title = data.get('abstractTitle', '').strip() or 'Без назви'
    section = data.get('sectionText', '').strip() or (f"Секція {data.get('targetSection')}" if data.get('targetSection') else 'Не вказана')
    format_text = data.get('partFormatText', '').strip() or data.get('partFormat', 'Не вказано')
    institution = data.get('institution', '').strip() or 'Не вказано'
    phone = data.get('phone', '').strip() or 'Не вказано'
    telegram = data.get('telegram', '').strip() or 'Не вказано'
    sub_id = data.get('submissionId', '') or f"USSF-{datetime.now().strftime('%Y%m%d-%H%M%S')}"

    # Read DOCX data
    docx_filename = os.path.basename(docx_path)
    try:
        with open(docx_path, 'rb') as f:
            docx_data = f.read()
    except Exception as e:
        return {"sent": False, "error": "FILE_READ_ERROR", "message": f"Не вдалося прочитати файл тез: {e}"}

    # Optional Program PDF attachment
    program_pdf_path = os.path.join(base_dir, cfg.get('program_pdf_path', 'USSF2026_Program.pdf'))
    program_data = None
    if os.path.isfile(program_pdf_path):
        try:
            with open(program_pdf_path, 'rb') as f:
                program_data = f.read()
        except Exception:
            program_data = None

    sent_results = []
    errors = []

    try:
        if smtp_port == 465:
            server = smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=20)
        else:
            server = smtplib.SMTP(smtp_host, smtp_port, timeout=20)
            server.ehlo()
            server.starttls()
            server.ehlo()

        server.login(smtp_user, smtp_pass)

        # 1. Committee Notification
        msg_comm = MIMEMultipart()
        msg_comm['From'] = f"USSF 2026 Registration <{smtp_user}>"
        msg_comm['To'] = committee_email
        msg_comm['Subject'] = f"[USSF 2026] Нова заявка {sub_id}: {author_initials} ({section})"

        comm_body = f"""Шановний Оргкомітет форуму USSF 2026!

Отримано нову реєстраційну заявку на I Всеукраїнський студентський хірургічний форум:

ID Заявки:         {sub_id}
ПІБ учасника:      {full_name} ({author_initials})
Email:             {author_email}
Телефон:           {phone}
Telegram:          {telegram}
Університет:       {institution}
Форма участі:      {format_text}
Секція:            {section}
Тема тез:          {title}

Науковий керівник: {data.get('scientificSupervisor', 'Не вказано')}
Кафедра:           {data.get('department', 'Не вказано')}
Завідувач кафедри: {data.get('headOfDepartment', 'Не вказано')}
Місто:             {data.get('cityCountry', 'Не вказано')}

Оформлені наукові тези за офіційним шаблоном НМУ у форматі Microsoft Word (.docx) додано до цього листа.

--
I Всеукраїнський студентський хірургічний форум (USSF 2026)
НМУ імені О.О. Богомольця, Київ
"""
        msg_comm.attach(MIMEText(comm_body, 'plain', 'utf-8'))

        # Attach DOCX
        att_comm = MIMEApplication(docx_data, _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document')
        att_comm.add_header('Content-Disposition', 'attachment', filename=('utf-8', '', docx_filename))
        msg_comm.attach(att_comm)

        server.send_message(msg_comm)
        sent_results.append(f"Сповіщення оргкомітету надіслано на {committee_email}")

        # 2. Participant Confirmation
        if author_email and cfg.get('send_copy_to_author', True):
            msg_auth = MIMEMultipart()
            msg_auth['From'] = f"Оргкомітет USSF 2026 <{smtp_user}>"
            msg_auth['To'] = author_email
            msg_auth['Subject'] = f"Ваша заявка на USSF 2026 зареєстрована (ID: {sub_id})"

            auth_body = f"""Шановний(а) {full_name}!

Щиро дякуємо за реєстрацію на I Всеукраїнський студентський хірургічний форум (USSF 2026) у Національному медичному університеті імені О.О. Богомольця!

Вашу заявку та матеріали успішно зафіксовано в базі наукового комітету.
Номер вашої заявки: {sub_id}

Деталі вашої участі:
• Форма участі: {format_text}
• Обрана секція: {section}
• Тема доповіді / тез: {title}

До цього листа ми прикріпили:
1. Оформлений документ ваших наукових тез за офіційним шаблоном НМУ у форматі Microsoft Word (.docx).
2. Офіційну програму форуму (USSF2026_Program.pdf).

Нагадуємо, що для студентів та інтернів участь є абсолютно безкоштовною. Після завершення форуму кожен зареєстрований учасник отримає офіційний іменний сертифікат.

З повагою,
Організаційний комітет USSF 2026
Національний медичний університет імені О.О. Богомольця
м. Київ, Україна
"""
            msg_auth.attach(MIMEText(auth_body, 'plain', 'utf-8'))

            # Attach DOCX
            att_auth = MIMEApplication(docx_data, _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document')
            att_auth.add_header('Content-Disposition', 'attachment', filename=('utf-8', '', docx_filename))
            msg_auth.attach(att_auth)

            # Attach Program PDF if available
            if program_data:
                att_prog = MIMEApplication(program_data, _subtype='pdf')
                att_prog.add_header('Content-Disposition', 'attachment', filename=('utf-8', '', 'USSF2026_Program.pdf'))
                msg_auth.attach(att_prog)

            server.send_message(msg_auth)
            sent_results.append(f"Підтвердження учаснику надіслано на {author_email}")

        server.quit()
        return {
            "sent": True,
            "status": "SENT",
            "recipients": [committee_email, author_email],
            "messages": sent_results,
            "filename": docx_filename
        }

    except Exception as e:
        return {
            "sent": False,
            "error": "SMTP_SEND_FAILED",
            "message": f"Помилка відправки через SMTP: {e}"
        }


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="USSF 2026 Academic Abstract DOCX Generator")
    parser.add_argument('--input', '-i', type=str, help="JSON file with submission data")
    parser.add_argument('--output', '-o', type=str, help="Output .docx file path")
    args = parser.parse_args()

    sample_data = {
        "fullName": "Кузнецов Олександр Олександрович",
        "abstractTitle": "Адаптивні методики підготовки пацієнтів до пластики великих післяопераційних вентральних гриж з втратою домену",
        "scientificSupervisor": "д.мед.н., професор Білянський Л. С.",
        "department": "Кафедра хірургії №1",
        "headOfDepartment": "д.мед.н., професор Пойда О. І.",
        "institution": "Національний медичний університет імені О. О. Богомольця",
        "cityCountry": "м. Київ, Україна",
        "abstractIntro": "Післяопераційна грижа — поширене ускладнення оперативних втручань на органах черевної порожнини [1].\n\nЗменшений об’єм черевної порожнини пов’язаний з хронічною латералізацією прямих м’язів живота.",
        "abstractAim": "Покращення результатів лікування хворих з великими дефектами черевної стінки з LOD.",
        "abstractMaterials": "Процедуру PPP проведено загалом у 45 пацієнтів.",
        "abstractResults": "Застосування всього комплексу представлених методик виконано у 18 пацієнтів.",
        "abstractConclusion": "Отримані результати підтверджують ефективність запропонованого алгоритму.",
        "abstractKeywords": "велика інцизійна грижа, втрата домену, пневмоперитонеум",
        "abstractReferences": "1. Smith, L., et al. (2024). Incisional Hernia Repair. JAWS, 3, 12452.\n2. Muysoms, F. E., et al. (2009). Classification of primary and incisional hernias. Hernia, 13(4), 407–414."
    }

    if args.input and os.path.isfile(args.input):
        with open(args.input, 'r', encoding='utf-8') as f:
            sample_data = json.load(f)

    out = create_abstract_docx(sample_data, args.output)
    print(f"DOCX created: {out}")
